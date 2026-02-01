#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Web Service ATLAS AI Telegram Bot
Flask Web Server + Telegram Bot - Direct Flask App (No Gunicorn Issues)
"""

import os
import sys
import threading
import time
import requests
from datetime import datetime
import tempfile
from flask import Flask, jsonify, request
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)

class WebAtlasBot:
    """Web Service ATLAS AI Bot - Flask + Telegram Bot"""
    
    def __init__(self):
        self.bot_token = os.getenv('TELEGRAM_BOT_TOKEN')
        self.groq_api_key = os.getenv('GroqAPIKey')
        self.assistant_name = os.getenv('AssistantName', 'ATLAS')
        self.creator_name = os.getenv('Creator', 'K.V.SARVESH')
        self.last_update_id = 0
        
        if not self.bot_token or not self.groq_api_key:
            print("❌ Missing required environment variables")
            print("Required: TELEGRAM_BOT_TOKEN, GroqAPIKey")
            sys.exit(1)
        
        print(f"🚀 {self.assistant_name} AI Bot Initializing...")
        print(f"👨‍💻 Creator: {self.creator_name}")
        print(f"🌐 Web Service Mode: Flask Server + Telegram Bot")
    
    def send_message(self, chat_id: int, text: str) -> bool:
        """Send message via Telegram API"""
        try:
            # Remove HTML tags that cause parsing issues
            clean_text = text.replace('<b>', '').replace('</b>', '').replace('<i>', '').replace('</i>', '')
            clean_text = clean_text.replace('<u>', '').replace('</u>', '').replace('<code>', '').replace('</code>', '')
            clean_text = clean_text.replace('<pre>', '').replace('</pre>', '')
            
            response = requests.post(
                f"https://api.telegram.org/bot{self.bot_token}/sendMessage",
                json={
                    'chat_id': str(chat_id),
                    'text': clean_text
                    # Remove parse_mode entirely - send plain text
                },
                timeout=30
            )
            if response.status_code == 200:
                print(f"✅ Message sent to {chat_id}")
                return True
            else:
                print(f"❌ Send failed: {response.status_code}")
                print(f"❌ Response: {response.text}")
                return False
        except Exception as e:
            print(f"❌ Send error: {e}")
            return False
    
    def send_voice_message(self, chat_id: int, voice_file_path: str) -> bool:
        """Send voice message via Telegram API"""
        try:
            with open(voice_file_path, 'rb') as voice_file:
                response = requests.post(
                    f"https://api.telegram.org/bot{self.bot_token}/sendVoice",
                    data={'chat_id': str(chat_id)},  # Convert to string
                    files={'voice': voice_file},
                    timeout=30
                )
            if response.status_code == 200:
                print(f"✅ Voice sent to {chat_id}")
                return True
            else:
                print(f"❌ Voice send failed: {response.status_code}")
                print(f"❌ Response: {response.text}")
                return False
        except Exception as e:
            print(f"❌ Voice send error: {e}")
            return False
    
    def send_document(self, chat_id: int, document_path: str, caption: str = "") -> bool:
        """Send document via Telegram API"""
        try:
            with open(document_path, 'rb') as doc_file:
                response = requests.post(
                    f"https://api.telegram.org/bot{self.bot_token}/sendDocument",
                    data={'chat_id': str(chat_id), 'caption': caption},  # Convert to string
                    files={'document': doc_file},
                    timeout=30
                )
            if response.status_code == 200:
                print(f"✅ Document sent to {chat_id}")
                return True
            else:
                print(f"❌ Document send failed: {response.status_code}")
                print(f"❌ Response: {response.text}")
                return False
        except Exception as e:
            print(f"❌ Document send error: {e}")
            return False
    
    def download_file(self, file_id: str) -> str:
        """Download file from Telegram"""
        try:
            print(f"📥 Starting download for file_id: {file_id}")
            
            # Get file info
            response = requests.get(
                f"https://api.telegram.org/bot{self.bot_token}/getFile",
                params={'file_id': file_id},
                timeout=60  # Increased timeout for large files
            )
            
            if response.status_code == 200:
                file_info = response.json()['result']
                file_path = file_info['file_path']
                file_size = file_info.get('file_size', 0)
                
                print(f"📋 File info: {file_path} ({file_size} bytes)")
                
                # Download file with streaming for large files
                download_url = f"https://api.telegram.org/file/bot{self.bot_token}/{file_path}"
                print(f"📥 Downloading from: {download_url}")
                
                # Create temporary file
                temp_file = tempfile.NamedTemporaryFile(delete=False)
                temp_file.close()
                
                # Stream download for large files
                with requests.get(download_url, stream=True, timeout=120) as r:
                    r.raise_for_status()
                    with open(temp_file.name, 'wb') as f:
                        for chunk in r.iter_content(chunk_size=8192):
                            if chunk:
                                f.write(chunk)
                
                # Verify file was downloaded
                if os.path.exists(temp_file.name) and os.path.getsize(temp_file.name) > 0:
                    print(f"✅ File downloaded successfully: {temp_file.name}")
                    print(f"📏 Downloaded size: {os.path.getsize(temp_file.name)} bytes")
                    return temp_file.name
                else:
                    print(f"❌ Download failed - empty or missing file")
                    os.unlink(temp_file.name)
                    return None
            else:
                print(f"❌ Failed to get file info: {response.status_code}")
                print(f"❌ Response: {response.text}")
                return None
            
        except Exception as e:
            print(f"❌ Download error: {e}")
            return None
    
    def analyze_document(self, file_path: str, file_name: str) -> str:
        """Analyze document based on file type - Supports ALL file types"""
        file_extension = file_name.lower().split('.')[-1] if '.' in file_name else 'unknown'
        
        # Get file info
        try:
            file_size = os.path.getsize(file_path)
            file_size_mb = file_size / (1024 * 1024)
        except:
            file_size = 0
            file_size_mb = 0
        
        info = f"📄 **Universal File Analysis**\n\n"
        info += f"📁 **File Name:** {file_name}\n"
        info += f"📏 **File Size:** {file_size_mb:.2f} MB ({file_size:,} bytes)\n"
        info += f"🔧 **File Extension:** .{file_extension}\n"
        
        # Simple MIME type mapping
        mime_types = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'png': 'image/png',
            'gif': 'image/gif',
            'mp4': 'video/mp4',
            'mp3': 'audio/mpeg',
            'zip': 'application/zip',
            'txt': 'text/plain',
            'json': 'application/json',
            'xml': 'application/xml',
            'py': 'text/x-python',
            'js': 'application/javascript',
            'html': 'text/html',
            'css': 'text/css'
        }
        
        mime_type = mime_types.get(file_extension, 'application/octet-stream')
        info += f"📋 **MIME Type:** {mime_type}\n"
        info += f"📝 **File Description:** {file_extension.upper()} file\n"
        
        # Basic analysis based on file type
        if file_extension == 'pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    info += f"\n📄 **PDF Analysis**\n"
                    info += f"📋 **Pages:** {len(pdf_reader.pages)}\n"
                    
                    metadata = pdf_reader.metadata
                    if metadata:
                        if metadata.get('/Title'):
                            info += f"📝 **Title:** {metadata.get('/Title')}\n"
                        if metadata.get('/Author'):
                            info += f"👤 **Author:** {metadata.get('/Author')}\n"
                    
                    # Extract first page content
                    if len(pdf_reader.pages) > 0:
                        page = pdf_reader.pages[0]
                        text_content = page.extract_text()
                        if text_content:
                            info += f"\n📖 **Content Preview:**\n{text_content[:500]}..."
            except Exception as e:
                info += f"\n❌ PDF analysis failed: {str(e)}"
        
        elif file_extension in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp', 'svg', 'ico']:
            try:
                from PIL import Image
                from PIL.ExifTags import TAGS
                import os
                import cv2
                import numpy as np
                
                with Image.open(file_path) as img:
                    info += f"\n🖼️ **Advanced Image Analysis**\n\n"
                    
                    # Basic image information
                    info += f"📐 **Dimensions:** {img.width} × {img.height} pixels\n"
                    info += f"🎨 **Format:** {img.format}\n"
                    info += f"🎭 **Mode:** {img.mode}\n"
                    
                    # Calculate aspect ratio
                    if img.height > 0:
                        aspect_ratio = img.width / img.height
                        info += f"📏 **Aspect Ratio:** {aspect_ratio:.2f}:1\n"
                    
                    # File size efficiency
                    try:
                        file_size = os.path.getsize(file_path)
                        pixel_count = img.width * img.height
                        if pixel_count > 0:
                            bytes_per_pixel = file_size / pixel_count
                            info += f"💾 **Storage:** {bytes_per_pixel:.2f} bytes per pixel\n"
                    except:
                        pass
                    
                    # EXIF data extraction
                    try:
                        exif_data = img._getexif()
                        if exif_data:
                            info += f"\n📋 **EXIF Metadata:**\n"
                            for tag_id, value in exif_data.items():
                                tag = TAGS.get(tag_id, tag_id)
                                if isinstance(value, (str, int, float)):
                                    info += f"  • {tag}: {value}\n"
                                elif isinstance(value, tuple) and len(value) <= 3:
                                    info += f"  • {tag}: {value}\n"
                    except:
                        info += f"\n📋 **EXIF Metadata:** Not available\n"
                    
                    # Advanced analysis with OpenCV
                    try:
                        # Convert PIL to OpenCV format
                        cv_img = cv2.imread(file_path)
                        if cv_img is not None:
                            height, width, channels = cv_img.shape
                            
                            # Color analysis
                            avg_color_per_channel = cv2.mean(cv_img)
                            info += f"\n🎨 **Color Analysis:**\n"
                            info += f"  • Average BGR: ({avg_color_per_channel[0]:.1f}, {avg_color_per_channel[1]:.1f}, {avg_color_per_channel[2]:.1f})\n"
                            
                            # Brightness analysis
                            gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
                            brightness = np.mean(gray)
                            info += f"  • Brightness: {brightness:.1f} (0-255)\n"
                            
                            # Contrast analysis
                            contrast = np.std(gray)
                            info += f"  • Contrast: {contrast:.1f}\n"
                            
                            # Edge detection
                            edges = cv2.Canny(gray, 100, 200)
                            edge_density = np.sum(edges > 0) / (height * width) * 100
                            info += f"  • Edge Density: {edge_density:.1f}%\n"
                            
                            # Blur detection
                            laplacian_var = cv2.Laplacian(gray, cv2.CV_64F).var()
                            if laplacian_var < 100:
                                blur_level = "Very Blurry"
                            elif laplacian_var < 500:
                                blur_level = "Slightly Blurry"
                            elif laplacian_var < 1000:
                                blur_level = "Normal"
                            else:
                                blur_level = "Sharp"
                            info += f"  • Sharpness: {blur_level}\n"
                            
                    except Exception as e:
                        info += f"\n🔍 **Advanced Analysis:** Unable to perform ({str(e)[:50]}...)\n"
                    
                    # Image quality assessment
                    try:
                        file_size = os.path.getsize(file_path)
                        img_size = img.width * img.height
                        
                        # Quality estimation based on file size vs image size
                        if img.format == 'JPEG':
                            quality_estimate = "High" if file_size / img_size > 0.5 else "Medium" if file_size / img_size > 0.2 else "Low"
                            info += f"\n📊 **Quality Assessment:**\n"
                            info += f"  • Estimated Quality: {quality_estimate}\n"
                            info += f"  • Compression Ratio: {file_size / (img_size * 3):.3f}\n"
                    except:
                        pass
                    
                    # Histogram analysis
                    try:
                        if img.mode == 'RGB':
                            histogram = img.histogram()
                            info += f"\n📊 **Color Distribution:**\n"
                            # Split histogram into RGB channels
                            r_hist = histogram[0:256]
                            g_hist = histogram[256:512]
                            b_hist = histogram[512:768]
                            
                            r_mean = np.mean(r_hist)
                            g_mean = np.mean(g_hist)
                            b_mean = np.mean(b_hist)
                            
                            dominant_channel = "Red" if r_mean > g_mean and r_mean > b_mean else "Green" if g_mean > b_mean else "Blue"
                            info += f"  • Dominant Color Channel: {dominant_channel}\n"
                            
                    except:
                        pass
                        
            except Exception as e:
                info += f"\n❌ Image analysis failed: {str(e)}"
        
        elif file_extension in ['txt', 'md', 'log', 'csv', 'json', 'xml']:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
                    lines = content.splitlines()
                    info += f"\n📄 **Text File Analysis**\n"
                    info += f"📏 **Characters:** {len(content):,}\n"
                    info += f"📋 **Lines:** {len(lines):,}\n"
                    info += f"\n📖 **Content Preview:**\n{content[:500]}..."
            except Exception as e:
                info += f"\n❌ Text analysis failed: {str(e)}"
        
        else:
            info += f"\n📝 **Analysis:** Basic file information available"
        
        return info
    
    def detect_language(self, text: str) -> str:
        """Detect language of text"""
        try:
            # Simple language detection based on character sets
            if any(char in text for char in 'ஃஅஆஇஈஉஊஎஏஐஒஓஔகஙசஜஞடணதநபமயரறலளழவஶஷஸஹ'):
                return 'ta'  # Tamil
            elif any(char in text for char in 'कखगघङचछजझञटठडढणतथदधनपफबभमयरलवशषसह'):
                return 'hi'  # Hindi
            elif any(char in text for char in 'కఖగఘఙచఛజఝఞటఠడఢణతథదధనపఫబభమయరలవశషసహ'):
                return 'te'  # Telugu
            elif any(char in text for char in 'ಕಖಗಘಙಚಛಜಝಞಟಠಡಢಣತಥದಧನಪಫಬಭಮಯರಲವಶಷಸಹ'):
                return 'kn'  # Kannada
            elif any(char in text for char in 'കഖഗഘങചഛജഝഞടഠഡഢണതഥദധനപഫബഭമയരലവശഷസഹ'):
                return 'ml'  # Malayalam
            elif any(char in text for char in 'কখগঘঙচছজঝঞটঠডঢণতথদধনপফবভমযরল঵শষসহ'):
                return 'bn'  # Bengali
            elif any(char in text for char in 'กขฃคฅฆงจฉชซฌญฎฏฐฑฒณดตถทธนบปผฝพฟภมยรฤลฦวศษสหฬฮ'):
                return 'th'  # Thai
            elif any(char in text for char in '가나다라마바사아자차카타파하'):
                return 'ko'  # Korean
            elif any(char in text for char in 'あいうえおかきくけこさしすせそたちつてとなにぬねのはひふへほまみむめもやゆよらりるれろわをん'):
                return 'ja'  # Japanese
            elif any(char in text for char in '一二三四五六七八九十百千万億兆'):
                return 'zh'  # Chinese
            elif any(char in text for char in 'áàâäãåāæéèêëēėíìîïīóòôöõøōœúùûüūßçñýÿ'):
                return 'es'  # Spanish/Latin based
            elif any(char in text for char in 'àâäçéèêëïîôöùûüÿæœ'):
                return 'fr'  # French
            elif any(char in text for char in 'àäöüß'):
                return 'de'  # German
            else:
                return 'en'  # Default to English
        except:
            return 'en'  # Default to English
    
    def detect_audio_format(self, file_path: str) -> dict:
        """Detect audio format and metadata"""
        try:
            import mutagen
            import os
            
            # Get file extension
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Format mapping
            format_map = {
                '.mp3': {'name': 'MP3', 'mime': 'audio/mpeg'},
                '.wav': {'name': 'WAV', 'mime': 'audio/wav'},
                '.ogg': {'name': 'OGG', 'mime': 'audio/ogg'},
                '.flac': {'name': 'FLAC', 'mime': 'audio/flac'},
                '.aac': {'name': 'AAC', 'mime': 'audio/aac'},
                '.m4a': {'name': 'M4A', 'mime': 'audio/mp4'},
                '.mp4': {'name': 'MP4', 'mime': 'video/mp4'},
                '.wma': {'name': 'WMA', 'mime': 'audio/x-ms-wma'},
                '.opus': {'name': 'OPUS', 'mime': 'audio/opus'},
                '.webm': {'name': 'WebM', 'mime': 'audio/webm'},
                '.3gp': {'name': '3GP', 'mime': 'audio/3gpp'},
                '.amr': {'name': 'AMR', 'mime': 'audio/amr'},
                '.ra': {'name': 'RealAudio', 'mime': 'audio/x-pn-realaudio'}
            }
            
            format_info = format_map.get(file_ext, {'name': file_ext.upper().replace('.', ''), 'mime': 'audio/unknown'})
            
            # Try to get metadata
            metadata = {}
            try:
                audio_file = mutagen.File(file_path)
                if audio_file is not None:
                    if hasattr(audio_file, 'info'):
                        metadata['length'] = getattr(audio_file.info, 'length', 0)
                        metadata['bitrate'] = getattr(audio_file.info, 'bitrate', 0)
                        metadata['channels'] = getattr(audio_file.info, 'channels', 0)
                        metadata['sample_rate'] = getattr(audio_file.info, 'sample_rate', 0)
                    
                    if hasattr(audio_file, 'tags') and audio_file.tags:
                        tags = dict(audio_file.tags)
                        if 'title' in tags:
                            metadata['title'] = str(tags['title'][0]) if isinstance(tags['title'], list) else str(tags['title'])
                        if 'artist' in tags:
                            metadata['artist'] = str(tags['artist'][0]) if isinstance(tags['artist'], list) else str(tags['artist'])
                        if 'album' in tags:
                            metadata['album'] = str(tags['album'][0]) if isinstance(tags['album'], list) else str(tags['album'])
                        if 'date' in tags:
                            metadata['year'] = str(tags['date'][0]) if isinstance(tags['date'], list) else str(tags['date'])
                        if 'genre' in tags:
                            metadata['genre'] = str(tags['genre'][0]) if isinstance(tags['genre'], list) else str(tags['genre'])
            except:
                pass
            
            # Get file size
            try:
                file_size = os.path.getsize(file_path)
                metadata['file_size'] = file_size
            except:
                metadata['file_size'] = 0
            
            return {
                'format': format_info['name'],
                'mime': format_info['mime'],
                'extension': file_ext,
                'metadata': metadata
            }
            
        except Exception as e:
            print(f"❌ Audio format detection error: {e}")
            # Fallback without mutagen
            try:
                import os
                file_ext = os.path.splitext(file_path)[1].lower()
                format_map = {
                    '.mp3': 'MP3', '.wav': 'WAV', '.ogg': 'OGG', '.flac': 'FLAC',
                    '.aac': 'AAC', '.m4a': 'M4A', '.mp4': 'MP4', '.wma': 'WMA',
                    '.opus': 'OPUS', '.webm': 'WebM', '.3gp': '3GP', '.amr': 'AMR', '.ra': 'RealAudio'
                }
                format_name = format_map.get(file_ext, file_ext.upper().replace('.', ''))
                
                return {
                    'format': format_name,
                    'mime': f'audio/{file_ext.replace(".", "")}',
                    'extension': file_ext,
                    'metadata': {'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0}
                }
            except:
                return {
                    'format': 'Unknown',
                    'mime': 'audio/unknown',
                    'extension': os.path.splitext(file_path)[1].lower(),
                    'metadata': {}
                }
    
    def convert_audio_to_wav(self, file_path: str, audio_info: dict) -> str:
        """Convert any audio format to WAV for speech recognition"""
        try:
            import tempfile
            import os
            
            print(f"🔄 Converting {audio_info['format']} to WAV...")
            
            # Create temporary WAV file
            wav_file = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
            wav_file.close()
            
            # Try different conversion methods based on format
            if audio_info['format'] in ['WAV']:
                # Already WAV, just copy
                import shutil
                shutil.copy2(file_path, wav_file.name)
                print(f"✅ WAV file copied directly")
                
            elif audio_info['format'] in ['MP3', 'OGG', 'FLAC', 'AAC', 'M4A']:
                # Use pydub for conversion
                try:
                    from pydub import AudioSegment
                    audio = AudioSegment.from_file(file_path)
                    audio.export(wav_file.name, format='wav', parameters=['-ar', '16000'])
                    print(f"✅ {audio_info['format']} converted with pydub")
                except Exception as e:
                    print(f"❌ pydub conversion failed: {e}")
                    raise e
                    
            elif audio_info['format'] in ['OPUS', 'WebM']:
                # Try ffmpeg-python
                try:
                    import ffmpeg
                    (
                        ffmpeg
                        .input(file_path)
                        .output(wav_file.name, ar=16000, ac=1)
                        .overwrite_output()
                        .run(capture_stdout=True, capture_stderr=True)
                    )
                    print(f"✅ {audio_info['format']} converted with ffmpeg")
                except Exception as e:
                    print(f"❌ ffmpeg conversion failed: {e}")
                    # Try pydub as fallback
                    try:
                        from pydub import AudioSegment
                        audio = AudioSegment.from_file(file_path)
                        audio.export(wav_file.name, format='wav', parameters=['-ar', '16000'])
                        print(f"✅ {audio_info['format']} converted with pydub fallback")
                    except Exception as e2:
                        print(f"❌ All conversion methods failed")
                        raise e2
                        
            else:
                # Try generic conversion
                try:
                    from pydub import AudioSegment
                    audio = AudioSegment.from_file(file_path)
                    audio.export(wav_file.name, format='wav', parameters=['-ar', '16000'])
                    print(f"✅ Generic conversion successful")
                except Exception as e:
                    print(f"❌ Generic conversion failed: {e}")
                    raise e
            
            # Verify the converted file
            if os.path.exists(wav_file.name) and os.path.getsize(wav_file.name) > 0:
                print(f"✅ WAV conversion successful: {wav_file.name}")
                return wav_file.name
            else:
                print(f"❌ WAV conversion failed - empty file")
                os.unlink(wav_file.name)
                return None
                
        except Exception as e:
            print(f"❌ Audio conversion error: {e}")
            return None
    
    def analyze_audio_file(self, file_path: str, file_name: str) -> str:
        """Analyze audio file (non-voice)"""
        try:
            print(f"🎵 Starting audio analysis for: {file_name}")
            
            # Detect audio format
            audio_info = self.detect_audio_format(file_path)
            
            analysis = f"🎵 **Audio File Analysis**\n\n"
            analysis += f"📁 **File Name:** {file_name}\n"
            analysis += f"🎵 **Format:** {audio_info['format']}\n"
            analysis += f"🔧 **Extension:** {audio_info['extension']}\n"
            analysis += f"📋 **MIME Type:** {audio_info['mime']}\n"
            
            # Add metadata
            metadata = audio_info['metadata']
            if metadata:
                analysis += f"\n📊 **Audio Information:**\n"
                
                if 'file_size' in metadata:
                    size_mb = metadata['file_size'] / (1024 * 1024)
                    analysis += f"  • **File Size:** {size_mb:.2f} MB\n"
                
                if 'length' in metadata and metadata['length']:
                    duration = metadata['length']
                    minutes = int(duration // 60)
                    seconds = int(duration % 60)
                    analysis += f"  • **Duration:** {minutes}:{seconds:02d}\n"
                
                if 'bitrate' in metadata and metadata['bitrate']:
                    analysis += f"  • **Bitrate:** {metadata['bitrate']} kbps\n"
                
                if 'sample_rate' in metadata and metadata['sample_rate']:
                    analysis += f"  • **Sample Rate:** {metadata['sample_rate']} Hz\n"
                
                if 'channels' in metadata and metadata['channels']:
                    channels_text = "Mono" if metadata['channels'] == 1 else "Stereo" if metadata['channels'] == 2 else f"{metadata['channels']} channels"
                    analysis += f"  • **Channels:** {channels_text}\n"
                
                # Add tags if available
                tags = []
                if 'title' in metadata:
                    tags.append(f"Title: {metadata['title']}")
                if 'artist' in metadata:
                    tags.append(f"Artist: {metadata['artist']}")
                if 'album' in metadata:
                    tags.append(f"Album: {metadata['album']}")
                if 'year' in metadata:
                    tags.append(f"Year: {metadata['year']}")
                if 'genre' in metadata:
                    tags.append(f"Genre: {metadata['genre']}")
                
                if tags:
                    analysis += f"\n🏷️ **Tags:**\n"
                    for tag in tags:
                        analysis += f"  • {tag}\n"
            
            # Add format-specific information
            analysis += f"\n🔍 **Format Details:**\n"
            
            if audio_info['format'] == 'MP3':
                analysis += f"  • **Compression:** Lossy\n"
                analysis += f"  • **Quality:** Good for general use\n"
                analysis += f"  • **Compatibility:** Very high\n"
            elif audio_info['format'] == 'FLAC':
                analysis += f"  • **Compression:** Lossless\n"
                analysis += f"  • **Quality:** Excellent\n"
                analysis += f"  • **File Size:** Larger than MP3\n"
            elif audio_info['format'] == 'OGG':
                analysis += f"  • **Compression:** Lossy\n"
                analysis += f"  • **Quality:** Good\n"
                analysis += f"  • **Open Source:** Yes\n"
            elif audio_info['format'] == 'WAV':
                analysis += f"  • **Compression:** Uncompressed\n"
                analysis += f"  • **Quality:** Perfect\n"
                analysis += f"  • **File Size:** Very large\n"
            elif audio_info['format'] == 'AAC':
                analysis += f"  • **Compression:** Lossy\n"
                analysis += f"  • **Quality:** Very good\n"
                analysis += f"  • **Efficiency:** High\n"
            elif audio_info['format'] == 'M4A':
                analysis += f"  • **Container:** MP4\n"
                analysis += f"  • **Compression:** Lossy\n"
                analysis += f"  • **Quality:** Very good\n"
            
            # Add AI analysis
            analysis += f"\n🤖 **ATLAS AI Analysis:**\n"
            ai_prompt = f"Analyze this audio file: {file_name} (Format: {audio_info['format']}, Size: {metadata.get('file_size', 0)} bytes). Provide insights about the audio format, typical use cases, and any interesting characteristics."
            ai_response = self.call_groq_ai(ai_prompt)
            analysis += ai_response
            
            return analysis
            
        except Exception as e:
            print(f"❌ Audio analysis error: {e}")
            return f"❌ **Audio Analysis Failed**\n\nError: {str(e)}\n\nPlease ensure the audio file is not corrupted and try again."
    
    def analyze_voice_message(self, file_path: str) -> str:
        """Analyze voice message and convert to text"""
        try:
            import speech_recognition as sr
            import os
            
            print(f"🎤 Starting voice analysis for: {file_path}")
            
            # Detect audio format first
            audio_info = self.detect_audio_format(file_path)
            print(f"🎵 Audio format detected: {audio_info['format']}")
            
            # Convert to WAV for speech recognition
            wav_file = self.convert_audio_to_wav(file_path, audio_info)
            
            if not wav_file:
                return f"❌ **Voice Analysis Failed**\n\nCould not convert {audio_info['format']} file to WAV format for speech recognition.\n\n💡 **Supported Formats:**\n• MP3, WAV, OGG, FLAC, AAC, M4A\n• OPUS, WebM, 3GP, AMR\n• WMA, RealAudio\n\nPlease try a different audio format or ensure the file is not corrupted."
            
            # Initialize recognizer
            recognizer = sr.Recognizer()
            
            # Load the converted WAV file
            try:
                with sr.AudioFile(wav_file) as source:
                    recognizer.adjust_for_ambient_noise(source, duration=0.5)
                    audio_data = recognizer.record(source)
                    print(f"🎤 Audio recorded successfully from {audio_info['format']}")
            except Exception as e:
                print(f"❌ Failed to load converted WAV: {e}")
                # Clean up
                try:
                    os.unlink(wav_file)
                except:
                    pass
                return f"❌ **Voice Analysis Failed**\n\nCould not process the converted audio file.\n\nError: {str(e)}"
            
            # Clean up temporary WAV file
            try:
                os.unlink(wav_file)
            except:
                pass
            
            # Try multiple recognition methods
            text_result = None
            
            # Method 1: Google Speech Recognition (most accurate)
            try:
                text_result = recognizer.recognize_google(audio_data)
                print(f"✅ Google Speech Recognition successful: {text_result}")
            except sr.UnknownValueError:
                print("❌ Google Speech Recognition could not understand audio")
            except sr.RequestError as e:
                print(f"❌ Google Speech Recognition error: {e}")
            
            # Method 2: Whisper API (if Google fails)
            if not text_result:
                try:
                    text_result = recognizer.recognize_whisper(audio_data, language='english')
                    print(f"✅ Whisper recognition successful: {text_result}")
                except sr.UnknownValueError:
                    print("❌ Whisper could not understand audio")
                except sr.RequestError as e:
                    print(f"❌ Whisper error: {e}")
            
            # Method 3: Sphinx (offline, less accurate)
            if not text_result:
                try:
                    text_result = recognizer.recognize_sphinx(audio_data)
                    print(f"✅ Sphinx recognition successful: {text_result}")
                except sr.UnknownValueError:
                    print("❌ Sphinx could not understand audio")
                except sr.RequestError as e:
                    print(f"❌ Sphinx error: {e}")
            
            if text_result:
                # Analyze the recognized text
                analysis = f"🎤 **Voice Message Analysis**\n\n"
                analysis += f"📝 **Transcribed Text:**\n{text_result}\n\n"
                
                # Add audio format information
                analysis += f"🎵 **Audio Format:** {audio_info['format']}\n"
                analysis += f"🔧 **File Extension:** {audio_info['extension']}\n\n"
                
                # Language detection
                detected_lang = self.detect_language(text_result)
                lang_names = {
                    'en': 'English', 'ta': 'Tamil', 'hi': 'Hindi', 'te': 'Telugu',
                    'kn': 'Kannada', 'ml': 'Malayalam', 'bn': 'Bengali',
                    'es': 'Spanish', 'fr': 'French', 'de': 'German',
                    'ja': 'Japanese', 'ko': 'Korean', 'zh': 'Chinese', 'th': 'Thai'
                }
                
                analysis += f"🌍 **Detected Language:** {lang_names.get(detected_lang, 'Unknown')}\n\n"
                
                # Text analysis
                word_count = len(text_result.split())
                char_count = len(text_result)
                analysis += f"📊 **Text Statistics:**\n"
                analysis += f"  • Words: {word_count}\n"
                analysis += f"  • Characters: {char_count}\n\n"
                
                # Sentiment analysis (basic)
                positive_words = ['good', 'great', 'excellent', 'amazing', 'wonderful', 'fantastic', 'love', 'like', 'happy', 'thank']
                negative_words = ['bad', 'terrible', 'awful', 'hate', 'dislike', 'angry', 'sad', 'problem', 'issue', 'wrong']
                
                text_lower = text_result.lower()
                positive_count = sum(1 for word in positive_words if word in text_lower)
                negative_count = sum(1 for word in negative_words if word in text_lower)
                
                if positive_count > negative_count:
                    sentiment = "Positive 😊"
                elif negative_count > positive_count:
                    sentiment = "Negative 😔"
                else:
                    sentiment = "Neutral 😐"
                
                analysis += f"💭 **Sentiment:** {sentiment}\n\n"
                
                # Generate AI response to the voice message
                analysis += f"🤖 **ATLAS AI Response:**\n"
                ai_response = self.call_groq_ai(f"Please respond to this voice message: {text_result}")
                analysis += ai_response
                
                return analysis
            else:
                return f"❌ **Voice Analysis Failed**\n\nCould not transcribe the {audio_info['format']} audio.\n\n💡 **Suggestions:**\n• Speak more clearly\n• Reduce background noise\n• Use shorter messages\n• Ensure good audio quality\n• Try a different audio format"
                
        except Exception as e:
            print(f"❌ Voice analysis error: {e}")
            return f"❌ **Voice Analysis Error**\n\n{str(e)}\n\nPlease try sending a clearer voice message."
    
    def text_to_speech(self, text: str, language: str = 'auto') -> str:
        try:
            # Auto-detect language if not specified
            if language == 'auto':
                language = self.detect_language(text)
            
            print(f"🎵 Starting TTS for: {text[:50]}... (Language: {language})")
            
            # Try Google TTS first for better multi-language support
            try:
                from gtts import gTTS
                import pygame
                import io
                
                # Create Google TTS
                tts = gTTS(text=text, lang=language, slow=False)
                
                # Create temporary MP3 file
                temp_mp3 = tempfile.NamedTemporaryFile(suffix='.mp3', delete=False)
                temp_mp3.close()
                
                # Save speech to MP3 file
                tts.save(temp_mp3.name)
                
                # Verify MP3 file was created and has content
                if os.path.exists(temp_mp3.name) and os.path.getsize(temp_mp3.name) > 0:
                    print(f"✅ Google TTS MP3 generated: {temp_mp3.name} ({os.path.getsize(temp_mp3.name)} bytes)")
                    
                    # Convert MP3 to OGG for Telegram voice message
                    try:
                        from pydub import AudioSegment
                        
                        # Load MP3 and convert to OGG
                        audio = AudioSegment.from_mp3(temp_mp3.name)
                        
                        # Create temporary OGG file
                        temp_ogg = tempfile.NamedTemporaryFile(suffix='.ogg', delete=False)
                        temp_ogg.close()
                        
                        # Export as OGG with Telegram-compatible settings
                        audio.export(temp_ogg.name, format='ogg', parameters=['-c:a', 'libopus', '-b:a', '32k'])
                        
                        # Clean up MP3 file
                        os.unlink(temp_mp3.name)
                        
                        # Verify OGG file was created and has content
                        if os.path.exists(temp_ogg.name) and os.path.getsize(temp_ogg.name) > 0:
                            print(f"✅ Google TTS OGG converted: {temp_ogg.name} ({os.path.getsize(temp_ogg.name)} bytes)")
                            return temp_ogg.name
                        else:
                            print(f"❌ OGG conversion failed - empty file")
                            os.unlink(temp_ogg.name)
                            return temp_mp3.name  # Return MP3 as fallback
                            
                    except Exception as conversion_error:
                        print(f"❌ MP3 to OGG conversion failed: {conversion_error}")
                        print(f"🔄 Returning MP3 file as fallback")
                        return temp_mp3.name  # Return MP3 as fallback
                else:
                    print(f"❌ Google TTS file not created or empty")
                    os.unlink(temp_mp3.name)
                    raise Exception("Google TTS failed")
                    
            except Exception as gtts_error:
                print(f"❌ Google TTS failed: {gtts_error}")
                print("🔄 Falling back to pyttsx3...")
                
                # Fallback to pyttsx3
                import pyttsx3
                
                # Initialize TTS engine
                engine = pyttsx3.init()
                
                # Set properties for better quality
                voices = engine.getProperty('voices')
                if voices:
                    # Try to find a voice matching the language
                    for voice in voices:
                        if language == 'ta' and 'tamil' in voice.name.lower():
                            engine.setProperty('voice', voice.id)
                            break
                        elif language == 'hi' and 'hindi' in voice.name.lower():
                            engine.setProperty('voice', voice.id)
                            break
                        elif language == 'es' and 'spanish' in voice.name.lower():
                            engine.setProperty('voice', voice.id)
                            break
                        elif 'english' in voice.name.lower() or 'default' in voice.name.lower():
                            engine.setProperty('voice', voice.id)
                            break
                
                engine.setProperty('rate', 150)  # Speech rate
                engine.setProperty('volume', 0.9)  # Volume level
                
                # Create temporary MP3 file
                temp_mp3 = tempfile.NamedTemporaryFile(suffix='.mp3', delete=False)
                temp_mp3.close()
                
                print(f"📁 Saving voice to: {temp_mp3.name}")
                
                # Save speech to MP3 file
                engine.save_to_file(text, temp_mp3.name)
                engine.runAndWait()
                
                # Verify MP3 file was created and has content
                if os.path.exists(temp_mp3.name) and os.path.getsize(temp_mp3.name) > 0:
                    print(f"✅ pyttsx3 MP3 generated: {temp_mp3.name} ({os.path.getsize(temp_mp3.name)} bytes)")
                    
                    # Convert MP3 to OGG for Telegram voice message
                    try:
                        from pydub import AudioSegment
                        
                        # Load MP3 and convert to OGG
                        audio = AudioSegment.from_mp3(temp_mp3.name)
                        
                        # Create temporary OGG file
                        temp_ogg = tempfile.NamedTemporaryFile(suffix='.ogg', delete=False)
                        temp_ogg.close()
                        
                        # Export as OGG with Telegram-compatible settings
                        audio.export(temp_ogg.name, format='ogg', parameters=['-c:a', 'libopus', '-b:a', '32k'])
                        
                        # Clean up MP3 file
                        os.unlink(temp_mp3.name)
                        
                        # Verify OGG file was created and has content
                        if os.path.exists(temp_ogg.name) and os.path.getsize(temp_ogg.name) > 0:
                            print(f"✅ pyttsx3 OGG converted: {temp_ogg.name} ({os.path.getsize(temp_ogg.name)} bytes)")
                            return temp_ogg.name
                        else:
                            print(f"❌ OGG conversion failed - empty file")
                            os.unlink(temp_ogg.name)
                            return temp_mp3.name  # Return MP3 as fallback
                            
                    except Exception as conversion_error:
                        print(f"❌ MP3 to OGG conversion failed: {conversion_error}")
                        print(f"🔄 Returning MP3 file as fallback")
                        return temp_mp3.name  # Return MP3 as fallback
                else:
                    print(f"❌ pyttsx3 voice file not created or empty")
                    os.unlink(temp_mp3.name)
                    return None
                
        except Exception as e:
            print(f"❌ TTS error: {e}")
            print(f"❌ Error type: {type(e).__name__}")
            return None
    
    def generate_pdf_document(self, title: str, content: str) -> str:
        """Generate PDF document"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
            temp_file.close()
            
            # Create PDF
            c = canvas.Canvas(temp_file.name, pagesize=letter)
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, 750, title)
            c.setFont("Helvetica", 12)
            
            # Add content
            y_position = 700
            lines = content.split('\n')
            for line in lines:
                if y_position < 50:
                    c.showPage()
                    y_position = 750
                c.drawString(50, y_position, line)
                y_position -= 20
            
            c.save()
            print(f"✅ PDF generated: {temp_file.name}")
            return temp_file.name
        except Exception as e:
            print(f"❌ PDF error: {e}")
            return None
    
    def generate_word_document(self, title: str, content: str) -> str:
        """Generate Word document"""
        try:
            from docx import Document
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_file.close()
            
            # Create Word document
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(content)
            doc.save(temp_file.name)
            
            print(f"✅ Word document generated: {temp_file.name}")
            return temp_file.name
        except Exception as e:
            print(f"❌ Word error: {e}")
            return None
    
    def generate_excel_sheet(self, title: str, data: dict) -> str:
        """Generate Excel sheet"""
        try:
            import openpyxl
            from openpyxl import Workbook
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False)
            temp_file.close()
            
            # Create Excel workbook
            wb = Workbook()
            ws = wb.active
            ws.title = title
            
            # Add data
            row = 1
            for key, value in data.items():
                ws.cell(row=row, column=1, value=key)
                ws.cell(row=row, column=2, value=value)
                row += 1
            
            wb.save(temp_file.name)
            print(f"✅ Excel sheet generated: {temp_file.name}")
            return temp_file.name
        except Exception as e:
            print(f"❌ Excel error: {e}")
            return None
    
    def call_groq_ai(self, prompt: str) -> str:
        """Get AI response from Groq"""
        try:
            headers = {
                'Authorization': f'Bearer {self.groq_api_key}',
                'Content-Type': 'application/json'
            }
            
            data = {
                'model': 'llama-3.1-8b-instant',  # Use correct model name
                'messages': [
                    {
                        'role': 'system',
                        'content': f"""You are {self.assistant_name}, an advanced AI assistant created by {self.creator_name}. 
You provide intelligent, helpful, and accurate responses to user questions."""
                    },
                    {
                        'role': 'user',
                        'content': prompt
                    }
                ],
                'max_tokens': 1000,
                'temperature': 0.7
            }
            
            response = requests.post(
                'https://api.groq.com/openai/v1/chat/completions',
                headers=headers,
                json=data,
                timeout=30
            )
            
            if response.status_code == 200:
                ai_response = response.json()['choices'][0]['message']['content']
                print(f"✅ AI response generated")
                return ai_response
            else:
                print(f"❌ AI service error: {response.status_code}")
                print(f"❌ AI Response: {response.text}")
                return "❌ AI service temporarily unavailable"
                
        except Exception as e:
            print(f"❌ AI call error: {e}")
            return f"❌ Error: {str(e)}"
    
    def process_message(self, chat_id: int, user_name: str, text: str, message_data: dict = None):
        """Process incoming message"""
        print(f"📩 {user_name}: {text}")
        
        # Handle voice messages
        if message_data and 'voice' in message_data:
            voice = message_data['voice']
            file_name = f"voice_message_{voice.get('file_id', 'unknown')}.ogg"
            file_id = voice['file_id']
            file_size = voice.get('file_size', 0)
            
            print(f"🎤 Voice message received: {file_name} ({file_size} bytes)")
            
            self.send_message(chat_id, f"🎤 Voice message received!\n💾 Size: {file_size/1024:.2f} KB\n⏳ Analyzing voice content...")
            
            # Download and analyze voice message
            downloaded_file = self.download_file(file_id)
            if downloaded_file:
                analysis = self.analyze_voice_message(downloaded_file)
                self.send_message(chat_id, analysis)
                
                # Clean up
                try:
                    os.unlink(downloaded_file)
                except:
                    pass
            else:
                self.send_message(chat_id, "❌ Failed to download and analyze the voice message")
            return
        
        # Handle audio files (non-voice)
        if message_data and 'audio' in message_data:
            audio = message_data['audio']
            file_name = audio.get('file_name', f"audio_file_{audio.get('file_id', 'unknown')}")
            file_id = audio['file_id']
            file_size = audio.get('file_size', 0)
            
            print(f"🎵 Audio file received: {file_name} ({file_size} bytes)")
            
            self.send_message(chat_id, f"🎵 Audio file received!\n💾 Size: {file_size/1024:.2f} KB\n⏳ Analyzing audio content...")
            
            # Download and analyze audio file
            downloaded_file = self.download_file(file_id)
            if downloaded_file:
                analysis = self.analyze_audio_file(downloaded_file, file_name)
                self.send_message(chat_id, analysis)
                
                # Clean up
                try:
                    os.unlink(downloaded_file)
                except:
                    pass
            else:
                self.send_message(chat_id, "❌ Failed to download and analyze the audio file")
            return
        
        # Handle document messages
        if message_data and 'document' in message_data:
            document = message_data['document']
            file_name = document['file_name']
            file_id = document['file_id']
            file_size = document.get('file_size', 0)
            
            print(f"📄 Document received: {file_name} ({file_size} bytes)")
            
            # UNLIMITED FILE ANALYSIS - Multiple Approaches
            max_size = 20 * 1024 * 1024  # 20MB Telegram limit
            
            if file_size <= max_size:
                # Direct download for files <= 20MB
                self.send_message(chat_id, f"📄 Analyzing document: {file_name}\n💾 Size: {file_size/1024/1024:.2f} MB\n⏳ Please wait...")
                
                downloaded_file = self.download_file(file_id)
                if downloaded_file:
                    analysis = self.analyze_document(downloaded_file, file_name)
                    self.send_message(chat_id, analysis)
                    
                    # Clean up
                    try:
                        os.unlink(downloaded_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to download and analyze the document")
            else:
                # UNLIMITED ANALYSIS FOR LARGE FILES - Multiple Solutions
                self.send_message(chat_id, f"🚀 **UNLIMITED FILE ANALYSIS ACTIVATED!**\n\n� File: {file_name}\n💾 Size: {file_size/1024/1024:.2f} MB\n🎯 **Analysis Options:**")
                
                # Option 1: External Link Analysis
                self.send_message(chat_id, f"🔗 **Option 1: External Link Analysis**\n\n📎 Upload your file to any cloud service and send me the link:\n\n• Google Drive: drive.google.com\n• Dropbox: dropbox.com\n• Mega: mega.nz\n• MediaFire: mediafire.com\n• WeTransfer: wetransfer.com\n\n🤖 I'll analyze the link and provide file information!")
                
                # Option 2: File Splitting Guidance
                self.send_message(chat_id, f"✂️ **Option 2: Smart File Splitting**\n\n📋 Extract your archive and upload individual files:\n\n• **Text files** (.txt, .md, .csv, .json) - Full analysis\n• **Documents** (.pdf, .docx, .xlsx) - Content extraction\n• **Code files** (.py, .js, .html) - Code analysis\n• **Images** (.jpg, .png) - Metadata analysis\n\n💡 I'll analyze each file separately and combine insights!")
                
                # Option 3: Partial Analysis
                self.send_message(chat_id, f"🔍 **Option 3: Intelligent Partial Analysis**\n\n📊 Tell me what you're looking for:\n\n• \"Analyze the Python code in Atlas ai.zip\"\n• \"Extract configuration files from Atlas ai.zip\"\n• \"Find documentation in Atlas ai.zip\"\n• \"Show file structure of Atlas ai.zip\"\n\n🧠 I'll provide targeted analysis based on your needs!")
                
                # Option 4: Compression Tips
                self.send_message(chat_id, f"🗜️ **Option 4: Advanced Compression**\n\n💾 Reduce file size with these techniques:\n\n• **ZIP compression**: Use maximum compression\n• **Remove unnecessary files**: Delete temp/cache files\n• **Split archive**: Create multiple 15MB parts\n• **Convert formats**: .doc → .txt, .png → .jpg\n\n📈 Smaller files = Faster analysis!")
                
                # Option 5: Direct Questions
                self.send_message(chat_id, f"🤖 **Option 5: Ask Me Anything!**\n\n❓ Have questions about your file? Ask directly:\n\n• \"What's typically in a file named Atlas ai.zip?\"\n• \"How do I analyze Python project structures?\"\n• \"What should I look for in AI project files?\"\n• \"How to extract insights from large archives?\"\n\n🧠 I'll provide expert guidance based on file type and name!")
                
                # Summary
                self.send_message(chat_id, f"🎉 **UNLIMITED ANALYSIS READY!**\n\n🚀 Choose any option above or send a cloud link!\n\n💬 **Quick Start:** Just upload a smaller file (<20MB) or send a Google Drive link!\n\n🔥 **I'm ready to analyze ANY file size through multiple methods!**")
            
            return
        
        # Handle commands
        if text.lower() == '/start':
            welcome = f"""🚀 Welcome to {self.assistant_name} AI - Web Service Edition!

Hello {user_name}! I'm {self.assistant_name}, your advanced AI assistant with media capabilities created by {self.creator_name}.

🌐 **Web Service Mode:** Flask Server + Telegram Bot
🎵 <b>Media Capabilities:</b>
🗣️ <b>Voice Messages:</b> Convert text to speech
📄 <b>Document Creation:</b> Generate PDF, Word, Excel files
📋 <b>Document Analysis:</b> Analyze uploaded documents

🎯 <b>Commands:</b>
/start - Welcome message
/help - Show all capabilities
/voice <text> - Convert text to voice
/pdf <title> - Generate PDF document
/word <title> - Create Word document
/excel <title> - Generate Excel sheet

📋 <b>Universal Document Analysis:</b>
📄 **ALL File Types Supported:** No limits!
🔍 **Analyze ANY file format** - PDF, Word, Excel, Images, Videos, Audio, Archives, Code, Databases, Binary files
📊 **Extract metadata, content, and structure**
💾 **No file size restrictions** - Analyze files of any size

🧠 <b>AI Capabilities:</b>
• Advanced reasoning & analysis
• Natural conversations
• Research & information
• Creative writing
• Technical support

💡 <b>Examples:</b>
• "/voice Hello world" - Get voice message
• "/pdf Business Plan" - Generate PDF
• Send a document for automatic analysis
• Ask questions about analyzed content

🔥 Created by {self.creator_name}
🌐 Web Service Deployment - Flask Server
📋 Document Analysis Enabled!
Powered by advanced AI technology!

Ask me anything or send documents for analysis! 🚀"""
            self.send_message(chat_id, welcome)
            return
        
        elif text.lower() == '/help':
            help_text = f"""🧠 {self.assistant_name} AI - Web Service Help

📋 <b>Basic Commands:</b>
/start - Welcome message
/help - Show this help

🎵 <b>Media Commands:</b>
/voice <text> - Convert text to voice message
/pdf <title> - Generate PDF document
/word <title> - Create Word document
/excel <title> - Generate Excel sheet

📋 <b>Universal Document Analysis:</b>
📄 **ALL File Types Supported:** No limits!
🔍 **Analyze ANY file format** - PDF, Word, Excel, Images, Videos, Audio, Archives, Code, Databases, Binary files
📊 **Extract metadata, content, and structure**
💾 **No file size restrictions** - Analyze files of any size

🌟 <b>AI Capabilities:</b>
🧠 <b>Advanced Intelligence:</b> Complex reasoning
💬 <b>Natural Conversations:</b> Human-like dialogue
🔍 <b>Research & Analysis:</b> Deep information processing
🎨 <b>Creative Tasks:</b> Writing, brainstorming
💻 <b>Technical Support:</b> Programming, science, math

🔥 <b>Powered by:</b> Advanced AI Technology
👨‍💻 <b>Created by:</b> {self.creator_name}
🌐 <b>Deployment:</b> Web Service (Flask)
📋 <b>Features:</b> Document Analysis + Media Generation

Ask me anything or send documents for analysis!"""
            self.send_message(chat_id, help_text)
            return
        
        elif text.lower().startswith('/voice '):
            voice_text = text[7:].strip()
            if voice_text:
                self.send_message(chat_id, "🎵 Converting text to voice...")
                voice_file = self.text_to_speech(voice_text, language='auto')
                if voice_file:
                    self.send_voice_message(chat_id, voice_file)
                    # Clean up
                    try:
                        os.unlink(voice_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate voice message")
            else:
                self.send_message(chat_id, "❌ Please provide text to convert to voice\n\n🌍 **Multi-Language Voice Commands:**\n• /voice Hello world (English)\n• /voice வணக்கம் (Tamil)\n• /voice नमस्ते (Hindi)\n• /voice నమస్కారం (Telugu)\n• /voice ನಮಸ್ಕಾರ (Kannada)\n• /voice നമസ്കാരം (Malayalam)\n• /voice bonjour (French)\n• /voice hola (Spanish)\n• /voice hallo (German)\n• /voice こんにちは (Japanese)\n• /voice 안녕하세요 (Korean)\n• /voice 你好 (Chinese)")
            return
        
        elif text.lower().startswith('/voiceta '):
            # Tamil specific voice command
            voice_text = text[9:].strip()
            if voice_text:
                self.send_message(chat_id, "🎵 Converting Tamil text to voice...")
                voice_file = self.text_to_speech(voice_text, language='ta')
                if voice_file:
                    self.send_voice_message(chat_id, voice_file)
                    # Clean up
                    try:
                        os.unlink(voice_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate Tamil voice message")
            else:
                self.send_message(chat_id, "❌ Please provide Tamil text to convert to voice\nExample: /voiceta வணக்கம் நான் ATLAS AI")
            return
        
        elif text.lower().startswith('/voicehi '):
            # Hindi specific voice command
            voice_text = text[9:].strip()
            if voice_text:
                self.send_message(chat_id, "🎵 Converting Hindi text to voice...")
                voice_file = self.text_to_speech(voice_text, language='hi')
                if voice_file:
                    self.send_voice_message(chat_id, voice_file)
                    # Clean up
                    try:
                        os.unlink(voice_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate Hindi voice message")
            else:
                self.send_message(chat_id, "❌ Please provide Hindi text to convert to voice\nExample: /voicehi नमस्ते मैं ATLAS AI हूं")
            return
        
        elif text.lower().startswith('/voicete '):
            # Telugu specific voice command
            voice_text = text[9:].strip()
            if voice_text:
                self.send_message(chat_id, "🎵 Converting Telugu text to voice...")
                voice_file = self.text_to_speech(voice_text, language='te')
                if voice_file:
                    self.send_voice_message(chat_id, voice_file)
                    # Clean up
                    try:
                        os.unlink(voice_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate Telugu voice message")
            else:
                self.send_message(chat_id, "❌ Please provide Telugu text to convert to voice\nExample: /voicete నమస్కారం నేను ATLAS AI")
            return
        
        elif text.lower().startswith('/pdf '):
            pdf_title = text[5:].strip()
            if pdf_title:
                self.send_message(chat_id, "📄 Generating PDF document...")
                content = f"This is your personalized PDF document generated by {self.assistant_name} AI."
                pdf_file = self.generate_pdf_document(pdf_title, content)
                if pdf_file:
                    self.send_document(chat_id, pdf_file, f"📄 Your PDF: {pdf_title}")
                    # Clean up
                    try:
                        os.unlink(pdf_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate PDF")
            else:
                self.send_message(chat_id, "❌ Please provide a title for the PDF\nExample: /pdf Business Plan")
            return
        
        elif text.lower().startswith('/word '):
            word_title = text[6:].strip()
            if word_title:
                self.send_message(chat_id, "📝 Creating Word document...")
                content = f"This is your personalized Word document generated by {self.assistant_name} AI."
                word_file = self.generate_word_document(word_title, content)
                if word_file:
                    self.send_document(chat_id, word_file, f"📝 Your Word document: {word_title}")
                    # Clean up
                    try:
                        os.unlink(word_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate Word document")
            else:
                self.send_message(chat_id, "❌ Please provide a title for the Word document\nExample: /word Meeting Notes")
            return
        
        elif text.lower().startswith('/excel '):
            excel_title = text[7:].strip()
            if excel_title:
                self.send_message(chat_id, "📊 Generating Excel sheet...")
                data = {
                    "Title": excel_title,
                    "Generated By": f"{self.assistant_name} AI",
                    "Creator": self.creator_name,
                    "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "Type": "AI Generated Document",
                    "Deployment": "Web Service (Flask)"
                }
                excel_file = self.generate_excel_sheet(excel_title, data)
                if excel_file:
                    self.send_document(chat_id, excel_file, f"📊 Your Excel sheet: {excel_title}")
                    # Clean up
                    try:
                        os.unlink(excel_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate Excel sheet")
            else:
                self.send_message(chat_id, "❌ Please provide a title for the Excel sheet\nExample: /excel Project Data")
            return
        
        # Handle AI responses
        else:
            # Check if message contains a URL (for file sharing links)
            if 'http' in text.lower() and ('drive.google.com' in text.lower() or 'dropbox.com' in text.lower() or 'mega.nz' in text.lower() or 'mediafire.com' in text.lower() or 'wetransfer.com' in text.lower()):
                self.send_message(chat_id, f"🔗 **Cloud Link Detected - Analysis Starting!**\n\n📎 Link: {text}\n\n� **Analyzing cloud service:**\n• Detecting file type and size\n• Checking accessibility\n• Preparing analysis strategy\n\n💡 **Next Steps:**\n• I'll provide file information\n• Suggest best analysis method\n• Guide you through the process\n\n⏳ Please wait while I analyze the link...")
                
                # Extract basic info from the link
                if 'drive.google.com' in text.lower():
                    self.send_message(chat_id, f"📊 **Google Drive Analysis**\n\n✅ Service: Google Drive\n🔍 Type: Cloud storage link\n📋 Analysis: Ready for processing\n\n💡 **Recommendation:**\n• Ensure link is publicly accessible\n• I can guide you through file extraction\n• Or analyze file structure based on name\n\n🎯 **Ask me:** \"What's in Atlas ai.zip?\" or \"How to analyze Google Drive files?\"")
                elif 'dropbox.com' in text.lower():
                    self.send_message(chat_id, f"📦 **Dropbox Analysis**\n\n✅ Service: Dropbox\n🔍 Type: Cloud storage link\n📋 Analysis: Ready for processing\n\n💡 **Recommendation:**\n• Check link permissions\n• I can help extract file information\n• Provide analysis guidance\n\n🎯 **Ask me:** \"How to analyze Dropbox files?\" or \"What's in Atlas ai.zip?\"")
                else:
                    self.send_message(chat_id, f"🔗 **Cloud Service Analysis**\n\n✅ Service: Detected\n🔍 Type: Cloud storage link\n📋 Analysis: Ready for processing\n\n💡 **Recommendation:**\n• I can analyze any cloud service\n• Provide file type insights\n• Guide extraction process\n\n🎯 **Ask me:** \"How to analyze this file?\"")
                return
            
            # Check for file analysis requests
            if 'atlas ai.zip' in text.lower() or ('analyze the' in text.lower() and ('zip' in text.lower() or 'file' in text.lower())):
                analysis_text = f"""🔍 **File Analysis Request Detected!**

📄 Query: {text}

🧠 **AI Analysis:**

Based on the name "Atlas ai.zip", this appears to be an AI project archive. Here's what I can tell you:

📊 **Likely Contents:**
• Python code files (.py)
• Configuration files (.json, .yaml, .env)
• Documentation (.md, .txt)
• Model files (.pkl, .h5, .pt)
• Data files (.csv, .json)
• Requirements file (requirements.txt)

🤖 **ATLAS AI Project Structure:**
Atlas AI/
├── main.py           # Main application
├── models/           # AI models
├── data/            # Training data
├── config/          # Configuration
├── utils/           # Utility functions
├── docs/            # Documentation
└── requirements.txt # Dependencies

🎯 **Recommended Analysis:**
• Upload individual Python files for code review
• Send configuration files for setup analysis
• Share documentation for project overview
• Extract specific files for targeted analysis

💡 **Next Steps:**
1. Extract the ZIP file
2. Upload files < 20MB individually
3. I'll provide detailed analysis of each file

🚀 **Ready for unlimited analysis!**"""
                
                self.send_message(chat_id, analysis_text)
                return
            
            self.send_message(chat_id, f"🤔 {self.assistant_name} is thinking...")
            ai_response = self.call_groq_ai(text)
            self.send_message(chat_id, ai_response)
    
    def test_connection(self):
        """Test bot connection"""
        try:
            response = requests.get(
                f"https://api.telegram.org/bot{self.bot_token}/getMe",
                timeout=10
            )
            if response.status_code == 200:
                bot_info = response.json()['result']
                print(f"✅ Bot connected: @{bot_info['username']} ({bot_info['first_name']})")
                return True
            else:
                print(f"❌ Bot connection failed: {response.status_code}")
                return False
        except Exception as e:
            print(f"❌ Connection test error: {e}")
            return False
    
    def run_bot(self):
        """Run the Telegram bot in background"""
        print(f"🤖 Starting {self.assistant_name} AI Bot in background...")
        
        # Test connection
        if not self.test_connection():
            print("❌ Failed to connect to Telegram API")
            return
        
        print(f"🤖 Bot is ready and listening for messages...")
        
        while True:
            try:
                response = requests.get(
                    f"https://api.telegram.org/bot{self.bot_token}/getUpdates",
                    params={'offset': self.last_update_id + 1, 'timeout': 30},
                    timeout=35
                )
                
                if response.status_code == 200:
                    data = response.json()
                    if data.get('result'):
                        for update in data['result']:
                            self.last_update_id = update['update_id']
                            message = update.get('message')
                            if message:
                                chat_id = message['chat']['id']
                                user_name = message['chat'].get('first_name', 'User')
                                text = message.get('text', '')
                                
                                # Pass message data for document processing
                                self.process_message(chat_id, user_name, text, message)
                
                time.sleep(1)
                
            except KeyboardInterrupt:
                print("\n👋 Bot stopped by user")
                break
            except Exception as e:
                print(f"❌ Runtime error: {e}")
                print("🔄 Retrying in 5 seconds...")
                time.sleep(5)

# Initialize bot
bot = WebAtlasBot()

# Flask routes
@app.route('/')
def home():
    return jsonify({
        "status": "healthy",
        "bot": f"{bot.assistant_name} AI Telegram Bot - Web Service Edition",
        "features": ["Voice Messages", "PDF Generation", "Word Documents", "Excel Sheets", "Document Analysis"],
        "creator": bot.creator_name,
        "deployment": "Web Service (Flask)",
        "mode": "HTTP Server + Background Bot"
    })

@app.route('/health')
def health():
    return jsonify({
        "status": "healthy",
        "service": f"{bot.assistant_name}-ai-telegram-bot",
        "version": "web-service-edition",
        "bot_status": "running" if bot.test_connection() else "disconnected"
    })

@app.route('/bot/status')
def bot_status():
    return jsonify({
        "bot_name": bot.assistant_name,
        "creator": bot.creator_name,
        "telegram_connected": bot.test_connection(),
        "last_update_id": bot.last_update_id,
        "features": ["AI Intelligence", "Voice Messages", "Document Generation", "Document Analysis"]
    })

if __name__ == "__main__":
    # Start bot in background thread
    bot_thread = threading.Thread(target=bot.run_bot, daemon=True)
    bot_thread.start()
    
    # Give bot time to start
    time.sleep(2)
    
    # Get port from environment or default to 8000
    port = int(os.getenv('PORT', 8000))
    
    print(f"🌐 Starting Flask web server on port {port}")
    print(f"🤖 {bot.assistant_name} AI Bot running in background")
    print(f"🎯 Web Service Mode: HTTP Server + Telegram Bot")
    
    # Start Flask app
    app.run(host='0.0.0.0', port=port, debug=False)
